import streamlit as st
import os
import io
import json
import re
import bcrypt
from datetime import datetime
import time

# --- Supabase Imports ---
from supabase import create_client, Client
from supabase.lib.client_options import ClientOptions
# from postgrest.exceptions import APIResponseException # Removed as per discussion to avoid ImportError

# --- AI & Document Processing Imports ---
from openai import OpenAI
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL

# --- Streamlit Page Configuration (MUST BE THE FIRST ST COMMAND) ---
# Set layout to wide to allow custom centering without Streamlit's default narrow column
st.set_page_config(
    page_title="SSO Consultants AI Recruitment",
    page_icon="🔍",
    layout="wide"
)

# --- Custom CSS for Styling ---
# This applies global styles to the Streamlit app to match the desired look.
st.markdown(
    """
    <style>
    /* Global base styling - pure white background, pure black text by default */
    body {
        background-color: #FFFFFF; /* Pure white background */
        color: #000000 !important; /* Pure black for general text readability - CRITICAL */
        font-family: 'Inter', sans-serif;
    }

    /* Hide Streamlit header and footer by default */
    .stApp > header {
        display: none;
    }
    .stApp > footer {
        display: none; /* We will render our own custom footer */
    }

    /* Main Streamlit app container and content blocks */
    .stApp, .css-18e3th9, .css-1d3f8gv {
        background-color: #FFFFFF; /* Ensure all main content areas are white */
        color: #000000 !important; /* Force black text for main content areas */
    }

    /* Specific targeting for ALL general text elements within the main Streamlit content area */
    /* This overrides any default Streamlit grey text */
    /* Targeting p, label, Streamlit-generated markdown/text spans, etc. */
    body p, body label, 
    .css-1d3f8gv p, .css-1d3f8gv label, 
    .logged-in-main-content p, .logged-in-main-content label,
    .stMarkdown span, .stText span, /* Targeting spans inside st.markdown/st.text where content actually resides */
    .stTextInput input[type="text"], .stTextInput input[type="password"], /* Input field text */
    .stTextInput label, .stFileUploader label,
    .stSelectbox label, .stRadio label,
    .stCheckbox label, .stDateInput label, .stNumberInput label, .stTextArea label,
    .stProgress, .stDataFrame {
        color: #000000 !important; /* Force all general text, labels, and alert text to pure black */
    }

    /* Customizing the sidebar - Now forcing to light gray */
    /* Targeting both the main sidebar container and its inner content area for robustness */
    .css-1lcbmhc, /* Main sidebar container */
    .css-1lcbmhc > section[data-testid="stSidebarContent"] { /* Inner content area */
        background-color: #F0F2F5 !important; /* FORCING Very Light Gray sidebar background - CRITICAL */
        color: #000000 !important; /* Default text in sidebar to black */
    }
    /* Sidebar text elements - Force to black now that background is light */
    .css-1lcbmhc .stRadio > label, 
    .css-1lcbmhc h1, .css-1lcbmhc h2, .css-1lcbmhc h3, .css-1lcbmhc h4, .css-1lcbmhc h5, .css-1lcbmhc h6, 
    .css-1lcbmhc p,
    .css-1lcbmhc .stMarkdown p, .css-1lcbmhc .stText p { /* Also target markdown/text within sidebar */
        color: #000000 !important; /* Force all sidebar text to pure black - CRITICAL */
    }
    /* Sidebar buttons - keep text white for contrast on blue background */
    .css-1lcbmhc .stButton > button {
        background-color: #0D47A1; 
        color: white; 
        border-radius: 0.5rem;
        border: none;
        padding: 0.5rem 1rem;
    }

    /* Styling for ALL buttons (st.button and st.form_submit_button) */
    .stButton > button, 
    .stForm button { /* Target buttons directly and buttons inside forms */
        background-color: #1976D2 !important; /* Vibrant blue from logo - CRITICAL for all buttons */
        color: white !important; /* Force button text to white - CRITICAL */
        border-radius: 0.5rem;
        padding: 0.75rem 1.5rem;
        font-size: 1.1rem;
        border: none;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        transition: all 0.3s ease;
        margin: 0.75rem; /* Consistent margin */
        min-width: 180px; /* Ensure buttons have a consistent minimum width */
    }
    .stButton > button:hover, 
    .stForm button:hover {
        background-color: #0D47A1 !important; /* Darker blue on hover - CRITICAL */
        transform: translateY(-2px);
    }

    /* Styling for forms and inputs */
    .stForm {
        padding: 2rem;
        border-radius: 0.75rem;
        background-color: #FFFFFF; /* White background for the form card */
        box-shadow: 0 8px 16px rgba(0, 0, 0, 0.1);
        margin-top: 2rem;
        width: 100%; /* Ensure form takes full width of its column */
        max-width: 500px; /* Limit form width for better appearance on large screens */
    }
    /* Labels within the form - explicitly pure black */
    .stForm .stTextInput > label, 
    .stForm .stSelectbox > label, 
    .stForm .stRadio > label,
    .stForm .stCheckbox > label {
        font-weight: bold;
        color: #000000 !important; /* Pure black for labels on white form background - CRITICAL */
        margin-bottom: 0.5rem;
    }
    /* Input fields (text typed by user) - explicitly pure black */
    .stForm .stTextInput input[type="text"], 
    .stForm .stTextInput input[type="password"] {
        color: #000000 !important; /* Pure black text for input fields - CRITICAL */
        background-color: #F8F8F8; /* Very light gray for input background */
        border-radius: 0.5rem;
        border: 1px solid #ced4da;
        padding: 0.75rem 1rem;
        width: 100%;
        margin-bottom: 1rem;
    }
    .stTextInput input:focus {
        border-color: #1976D2; /* Focus color matching logo blue */
        box-shadow: 0 0 0 0.2rem rgba(25, 118, 210, 0.25); 
    }
    
    /* Styling for Radio Buttons (User/Admin under Assign Role AND Sidebar Navigation Radio Buttons) */
    /* Target the text itself inside the radio options */
    .stRadio div[data-testid="stRadio"] label span p {
        color: #000000 !important; /* Pure black for 'User', 'Admin', 'Dashboard' etc. text - CRITICAL */
    }
    /* Style the radio button circles (unselected) */
    .stRadio div[data-testid="stRadio"] input[type="radio"] + div::before {
        background-color: #FFFFFF !important; /* White background for unselected */
        border: 2px solid #333333 !important; /* Dark border for unselected */
        width: 18px !important; /* Consistent size */
        height: 18px !important; /* Consistent size */
        top: 3px !important; /* Adjust vertical alignment */
        left: 0px !important; /* Adjust horizontal alignment */
    }
    /* Style the radio button circles (selected dot) */
    .stRadio div[data-testid="stRadio"] input[type="radio"]:checked + div::after {
        background-color: #1976D2 !important; /* Vibrant Blue for selected dot - CRITICAL */
        width: 10px !important; /* Size of the inner dot */
        height: 10px !important; /* Size of the inner dot */
        top: 7px !important; /* Adjust vertical alignment of dot */
        left: 4px !important; /* Adjust horizontal alignment of dot */
    }

    /* Success/Error/Warning/Info messages */
    .stAlert {
        border-radius: 0.5rem;
        margin-top: 1rem;
        margin-bottom: 1rem;
    }
    .stAlert.success {
        background-color: #d4edda; /* Light green */
        color: #155724 !important; /* Dark green text - CRITICAL */
        border-color: #c3e6cb;
    }
    .stAlert.error {
        background-color: #f8d7da; /* Light red */
        color: #721c24 !important; /* Dark red text - CRITICAL */
        border-color: #f5c6cb;
    }
    .stAlert.warning {
        background-color: #fff3cd; /* Light yellow */
        color: #856404 !important; /* Dark yellow text - CRITICAL */
        border-color: #ffeeba;
    }
    .stAlert.info { /* Explicitly targeting info alerts */
        background-color: #d1ecf1; /* Light blue info box background */
        color: #000000 !important; /* Pure black text - CRITICAL */
        border-color: #bee5eb;
    }


    /* Specific style for the initial info message on login page */
    .initial-info-message {
        font-size: 1.1em;
        color: #000000 !important; /* Pure black for clear visibility - CRITICAL */
        margin-top: 1.5rem; 
        margin-bottom: 2rem;
        font-style: italic;
    }

    /* Centering content within a column (applied to Streamlit's main block) */
    .st-emotion-cache-16txt4v { 
        display: flex;
        flex-direction: column;
        align-items: center; /* Center horizontally */
        justify-content: flex-start; /* Align from top vertically */
        min-height: 90vh; /* Ensure content pushes footer down on shorter pages */
        padding-top: 3.5rem; /* Adjusted padding from top for the main title */
    }

    /* Styling for the central main title */
    .main-app-title {
        color: #0D47A1 !important; /* Deep Dark Blue for main title - CRITICAL */
        font-size: 2.8em; 
        font-weight: bold;
        margin-bottom: 0.5rem; 
        text-align: center; /* Explicitly center align */
    }
    .sub-app-title {
        color: #0D47A1 !important; /* Deep Dark Blue for subtitle - CRITICAL */
        font-size: 1.3em; 
        margin-bottom: 2.5rem; 
        text-align: center; /* Explicitly center align */
    }

    /* Specific targeting for all h1, h2, h3, h4, h5, h6 tags */
    h1, h2, h3, h4, h5, h6 {
        color: #0D47A1 !important; /* Deep Dark Blue for all headings - CRITICAL */
    }
    /* Override for the login form h3 to be pure black as requested */
    /* Target the specific generated Streamlit h3 elements for the login prompt */
    .st-emotion-cache-nahz7x h3, .st-emotion-cache-nahz7x { 
        color: #000000 !important; /* Pure black for the login mode title - CRITICAL */
    }


    /* Top-right logo container */
    .top-right-logo {
        position: fixed; 
        top: 10px; 
        right: 10px; 
        z-index: 9999; 
        background-color: rgba(255, 255, 255, 0.0); /* Transparent background */
        padding: 5px;
        border-radius: 8px;
        border: 2px solid red; /* *** DEBUGGING BORDER - RETAINED AS REQUESTED *** */
    }
    .top-right-logo img {
        width: 100px; 
        height: auto;
    }

    /* Adjust padding/alignment for logged in pages, overriding centering for content */
    .logged-in-main-content .st-emotion-cache-16txt4v {
        align-items: flex-start; /* Reset to left align */
        padding-top: 2rem; 
        text-align: left;
        margin-left: 1rem; 
        margin-right: 1rem; 
        width: calc(100% - 2rem); 
    }
    /* Force all text elements within the main content area (after login) to be pure black */
    /* This is a broad rule for safety */
    .logged-in-main-content p, 
    .logged-in-main-content .stMarkdown, 
    .logged-in-main-content .stText, 
    .logged-in-main-content .stInfo, 
    .logged-in-main-content .stWarning,
    .logged-in-main-content .stError,
    .logged-in-main-content label,
    .logged-in-main-content .stSelectbox,
    .logged-in-main-content .stRadio,
    .logged-in-main-content .stCheckbox,
    .logged-in-main-content .stDateInput,
    .logged-in-main-content .stNumberInput,
    .logged-in-main-content .stTextArea,
    .logged-in-main-content .stProgress,
    .logged-in-main-content .stDataFrame {
        color: #000000 !important; /* Pure black for all general text and labels - CRITICAL */
    }
    /* Ensure headings on logged-in pages are deep dark blue */
    .logged-in-main-content h1, 
    .logged-in-main-content h2, 
    .logged-in-main-content h3,
    .logged-in-main-content h4,
    .logged-in-main-content h5,
    .logged-in-main-content h6 {
        text-align: left; 
        color: #0D47A1 !important; /* Deep Dark Blue for headings when logged in - CRITICAL */
    }
    .logged-in-main-content .stForm {
        width: auto; 
        max-width: none; 
    }


    /* Hide the default Streamlit hamburger menu button and Share button */
    .css-hi6a2p { 
        display: none !important;
    }
    .css-1dp5x4b { 
        display: none !important;
    }
    .css-1gh6j8x { 
        display: none !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# --- Inject Top-Right Logo HTML ---
st.markdown(
    f"""
    <div class="top-right-logo">
        <img src="https://raw.githubusercontent.com/SsoConsultations/sso-consultants-recruitment-app/main/logo.png" alt="Company Logo" onerror="this.onerror=null; this.src='https://placehold.co/100x100/A0A0A0/FFFFFF?text=Logo+Missing';">
    </div>
    """,
    unsafe_allow_html=True
)

# --- Configuration: NOW READING FROM ENVIRONMENT VARIABLES ---
SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY")
SUPABASE_SERVICE_ROLE_KEY = os.environ.get("SUPABASE_SERVICE_ROLE_KEY") # ADDED FOR ADMIN OPERATIONS

# --- Supabase Initialization Function ---
def initialize_supabase_app():
    """
    Initializes the Supabase client and stores it in session state.
    This function is called only once per app run or when 'supabase_client' is not in session state.
    """
    print("DEBUG: Attempting to initialize Supabase client...")
    try:
        if not SUPABASE_URL or not SUPABASE_KEY or not SUPABASE_SERVICE_ROLE_KEY: # MODIFIED: Check for service key too
            print("ERROR: Supabase URL, Key, or Service Role Key not found in environment variables.")
            st.error("Supabase URL, Key, or Service Role Key not found in environment variables. Please configure them.")
            st.stop()

        supabase_client_instance = create_client(SUPABASE_URL, SUPABASE_KEY, options=ClientOptions(postgrest_client_timeout=10))
        st.session_state['supabase_client'] = supabase_client_instance
        print("DEBUG: Supabase client initialized successfully and stored in session state.")
        print(f"DEBUG: Session state 'supabase_client' is now: {type(st.session_state['supabase_client'])}")

        # ADDED: Initialize service role client
        supabase_service_role_client_instance = create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY, options=ClientOptions(postgrest_client_timeout=10))
        st.session_state['supabase_service_role_client'] = supabase_service_role_client_instance
        print("DEBUG: Supabase service role client initialized successfully and stored in session state.")
        print(f"DEBUG: Session state 'supabase_service_role_client' is now: {type(st.session_state['supabase_service_role_client'])}")

    except Exception as e:
        print(f"ERROR: Error during Supabase initialization: {e}")
        st.error(f"Error initializing Supabase: {e}. Please ensure your environment variables are correctly configured.")
        st.stop()

# --- Ensure Supabase is initialized and client is available in session state ---
if 'supabase_client' not in st.session_state or st.session_state['supabase_client'] is None:
    print("DEBUG: 'supabase_client' not found in session state or is None. Calling initialize_supabase_app().")
    initialize_supabase_app()
else:
    print("DEBUG: 'supabase_client' already exists in session state. Supabase previously initialized.")

supabase = st.session_state['supabase_client']
# Note: st.session_state['supabase_service_role_client'] is available directly via session state where needed.

# --- Initialize OpenAI client ---
try:
    OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
    if not OPENAI_API_KEY:
        st.error("OpenAI API key not found in environment variables. Please configure it.")
        st.stop()
    openai_client = OpenAI(api_key=OPENAI_API_KEY)
    print("DEBUG: OpenAI client initialized successfully.")
except Exception as e:
    st.error(f"OpenAI client not initialized: {e}. Please check your OPENAI_API_KEY in environment variables.")
    print(f"ERROR: OpenAI client initialization failed: {e}")
    st.stop()

# Admin Credentials (for a hardcoded admin user, outside Supabase Auth)
ADMIN_EMAIL = os.environ.get("ADMIN_EMAIL", "admin@sso.com")
# Default hash for "adminpass" if not set, for initial setup convenience.
# In production, this should always be set via environment variable with a strong, generated hash.
ADMIN_PASSWORD_HASH = os.environ.get("ADMIN_PASSWORD_HASH", bcrypt.hashpw("adminpass".encode('utf-8'), bcrypt.gensalt()).decode('utf-8')) # MODIFIED: Default hash for "adminpass" if not set for initial setup


# --- Streamlit Session State Initialization ---
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
if 'user_name' not in st.session_state:
    st.session_state['user_name'] = ''
if 'user_email' not in st.session_state:
    st.session_state['user_email'] = ''
if 'user_uid' not in st.session_state:
    st.session_state['user_uid'] = ''
if 'is_admin' not in st.session_state:
    st.session_state['is_admin'] = False
if 'ai_review_result' not in st.session_state:
    st.session_state['ai_review_result'] = None
if 'generated_docx_buffer' not in st.session_state:
    st.session_state['generated_docx_buffer'] = None
if 'review_triggered' not in st.session_state:
    st.session_state['review_triggered'] = False
if 'current_page' not in st.session_state:
    st.session_state['current_page'] = 'Login'
if 'jd_filename_for_save' not in st.session_state:
    st.session_state['jd_filename_for_save'] = "Job Description"
if 'cv_filenames_for_save' not in st.session_state:
    st.session_state['cv_filenames_for_save'] = []
if 'login_mode' not in st.session_state:
    st.session_state['login_mode'] = None
if 'new_user_email_for_pw_reset' not in st.session_state:
    st.session_state['new_user_email_for_pw_reset'] = ''
if 'new_user_uid_for_pw_reset' not in st.session_state:
    st.session_state['new_user_uid_for_pw_reset'] = ''

# --- Helper Functions for Text Extraction ---
def extract_text_from_pdf(uploaded_file_bytes_io):
    """Extracts text from a PDF file using PyPDF2."""
    try:
        reader = PdfReader(uploaded_file_bytes_io)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        print(f"ERROR (extract_text_from_pdf): {e}")
        return None

def extract_text_from_docx(uploaded_file_bytes_io):
    """Extracts text from a DOCX file using python-docx."""
    try:
        document = Document(uploaded_file_bytes_io)
        text = ""
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        print(f"ERROR (extract_text_from_docx): {e}")
        return None

def get_file_content(uploaded_file_bytes_io, filename):
    """Determines file type based on extension and extracts text content."""
    file_extension = os.path.splitext(filename)[1].lower()

    if file_extension == '.pdf':
        return extract_text_from_pdf(uploaded_file_bytes_io)
    elif file_extension == '.docx':
        return extract_text_from_docx(uploaded_file_bytes_io)
    elif file_extension == '.txt':
        return uploaded_file_bytes_io.read().decode('utf-8')
    else:
        st.error(f"Unsupported file type: {file_extension}. Only PDF, DOCX, TXT are supported.")
        print(f"ERROR (get_file_content): Unsupported file type {file_extension} for {filename}")
        return None

# --- AI Function: Comparative Analysis ---
def get_comparative_ai_analysis(jd_text, all_cv_data):
    """
    Uses OpenAI to perform a comparative analysis of multiple CVs against a JD.
    Returns a complex JSON object containing a table of candidate evaluations,
    a table of criteria observations, additional observations text, and a final
    shortlist recommendation.
    """
    if not jd_text or not all_cv_data:
        print("DEBUG (get_comparative_ai_analysis): Missing JD or CV data.")
        return {"error": "Missing Job Description or Candidate CV content for comparative analysis."}

    # System prompt defines the AI's role and the required JSON output format
    system_prompt = """
    You are an expert Talent Acquisition professional in India. Your task is to perform a detailed comparative analysis of multiple candidate CVs against a given Job Description (JD).

    Your output MUST be a single JSON object with the following structure, containing two distinct arrays for tables and two strings for text sections:
    {
      "candidate_evaluations": [
        {
          "Candidate Name": "...",       // Derived from filename (e.g., "Charlie")
          "Match %": "...",              // Numerical percentage as a string (e.g., "85%")
          "Ranking": "...",              // E.g., "1", "2", "3", etc. (NO MEDALS)
          "Shortlist Probability": "...",// E.g., "High", "Moderate", "Low"
          "Key Strengths": "...",        // Concise points, comma-separated or short phrase. Highlight relevant relevant experience.
          "Key Gaps": "...",             // Concise points, comma-separated or short phrase.
          "Location Suitability": "...", // E.g., "Pune", "Delhi (flexible)", "Remote", "Not Specified"
          "Comments": "..."              // Any other relevant observation for this candidate, including fit for Indian context.
        },
        // ... more candidate evaluation objects for each CV ...
      ],
      "criteria_observations": [ // This array is for the second table comparing candidates across common criteria
        {
          "Criteria": "Education (MBA HR)",
          "Candidate 1 Name": "✅/❌/⚠️", // Column for each candidate provided (e.g., "Himanshukulkarni")
          "Candidate 2 Name": "✅/❌/⚠️",
          // ... more candidate columns based on actual input filenames
        },
        {
          "Criteria": "Recruitment / TA Experience",
          "Candidate 1 Name": "✅/❌/⚠️",
          "Candidate 2 Name": "✅/❌/⚠️",
        },
        // ... more criteria rows ...
      ],
      "additional_observations_text": "...", // Comprehensive text for general observations not covered in tables.
      "final_shortlist_recommendation": "..." // Concise text for the final recommendation, explicitly naming shortlisted candidates.
    }

    Ensure "Match %" is a string.
    Ensure "Ranking" is a numerical rank string (e.g., "1", "2") without any emoji symbols (like 🥇, 🥈, 🥉).
    For "criteria_observations", dynamically create columns for each candidate using their names (e.g., "Gauri Deshmukh", "Himanshukulkarni"). Use ✅ for good fit, ❌ for not a fit, ⚠️ for partial fit.
    Make sure all text fields are within the string limits of JSON.
    The "Candidate Name" in "candidate_evaluations" and the dynamic column headers in "criteria_observations" should be derived from the provided filenames (e.g., "Gauri CV.pdf" -> "Gauri").
    """

    user_prompt = f"""
    Here is the Job Description (JD):
    ---
    {jd_text}
    ---

    Here are the Candidate CVs for comparative analysis:
    """
    # Append each CV's content to the user prompt
    for idx, cv_item in enumerate(all_cv_data):
        # Extract name without extension for table headers
        candidate_name_for_prompt = os.path.splitext(cv_item['filename'])[0].replace(" CV", "").strip()
        user_prompt += f"\n--- Candidate {idx+1} (Name: {candidate_name_for_prompt}, Filename: {cv_item['filename']}) ---\n"
        user_prompt += f"{cv_item['text']}\n"
    user_prompt += "--- End of Candidate CVs ---"
    user_prompt += "\n\nPlease provide the comparative analysis in the specified JSON format."

    try:
        with st.spinner("AI is analyzing the JD and CVs... This may take a moment."):
            print("DEBUG (get_comparative_ai_analysis): Sending request to OpenAI API.")
            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.2,
                response_format={"type": "json_object"}
            )
        ai_response_content = response.choices[0].message.content
        print(f"DEBUG (get_comparative_ai_analysis): Raw AI Response: {ai_response_content[:200]}...")
        comparative_data = json.loads(ai_response_content)

        if "candidate_evaluations" in comparative_data:
            for candidate in comparative_data["candidate_evaluations"]:
                if "Ranking" in candidate and isinstance(candidate["Ranking"], str):
                    candidate["Ranking"] = re.sub(r'[\U0001F3C5-\U0001F3CA\U0001F947-\U0001F949]', '', candidate["Ranking"]).strip()

        print("DEBUG (get_comparative_ai_analysis): AI analysis successful.")
        return comparative_data

    except json.JSONDecodeError as e:
        st.error(f"Error: AI response was not valid JSON. Please try again or refine input. Error: {e}")
        st.code(ai_response_content)
        print(f"ERROR (get_comparative_ai_analysis): JSON Decode Error: {e}, Response: {ai_response_content}")
        return {"error": f"AI response format error: {e}"}
    except Exception as e:
        st.error(f"An unexpected error occurred during AI analysis: {e}")
        print(f"ERROR (get_comparative_ai_analysis): Unexpected Error during AI analysis: {e}")
        return {"error": f"AI processing failed: {e}"}

# --- DOCX Generation Function ---
def generate_docx_report(comparative_data, jd_filename="Job Description", cv_filenames_str="Candidates"):
    """
    Generates a DOCX report based on the comparative AI analysis data.
    Includes two tables and text sections.
    """
    try:
        document = Document()

        section = document.sections[0]
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        document.add_heading("JD-CV Comparative Analysis Report", level=0)
        document.add_paragraph().add_run("Generated by SSO Consultants AI").italic = True
        document.add_paragraph().add_run(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").small_caps = True
        document.add_paragraph(f"Job Description: {jd_filename}\nCandidates: {cv_filenames_str}")
        document.add_paragraph("\n")

        candidate_evaluations_data = comparative_data.get("candidate_evaluations", [])
        criteria_observations_data = comparative_data.get("criteria_observations", [])
        additional_observations_text = comparative_data.get("additional_observations_text", "No general observations provided.")
        final_shortlist_recommendation = comparative_data.get("final_shortlist_recommendation", "No final recommendation provided.")

        if candidate_evaluations_data:
            document.add_heading("🧾 Candidate Evaluation Table", level=1)
            document.add_paragraph("Detailed assessment of each candidate against the Job Description:")

            df_evaluations = pd.DataFrame(candidate_evaluations_data)

            expected_cols_eval = ["Candidate Name", "Match %", "Ranking", "Shortlist Probability", "Key Strengths", "Key Gaps", "Location Suitability", "Comments"]
            for col in expected_cols_eval:
                if col not in df_evaluations.columns:
                    df_evaluations[col] = "N/A"
            df_evaluations = df_evaluations[expected_cols_eval]

            table_eval = document.add_table(rows=1, cols=len(df_evaluations.columns))
            table_eval.style = 'Table Grid'

            hdr_cells_eval = table_eval.rows[0].cells
            for i, col_name in enumerate(df_evaluations.columns):
                hdr_cells_eval[i].text = col_name
                for paragraph in hdr_cells_eval[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)
                hdr_cells_eval[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            for index, row in df_evaluations.iterrows():
                row_cells = table_eval.add_row().cells
                for i, cell_value in enumerate(row):
                    row_cells[i].text = str(cell_value)
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

            document.add_paragraph("\n")

        if criteria_observations_data:
            document.add_heading("✅ Additional Observations (Criteria Comparison)", level=1)

            df_criteria = pd.DataFrame(criteria_observations_data)

            table_criteria = document.add_table(rows=1, cols=len(df_criteria.columns))
            table_criteria.style = 'Table Grid'

            hdr_cells_criteria = table_criteria.rows[0].cells
            for i, col_name in enumerate(df_criteria.columns):
                hdr_cells_criteria[i].text = col_name
                for paragraph in hdr_cells_criteria[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)
                hdr_cells_criteria[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            for index, row in df_criteria.iterrows():
                row_cells = table_criteria.add_row().cells
                for i, cell_value in enumerate(row):
                    row_cells[i].text = str(cell_value)
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

            document.add_paragraph("\n")

        if additional_observations_text and additional_observations_text.strip() not in ["No general observations provided.", ""]:
            document.add_heading("General Observations", level=2)
            document.add_paragraph(additional_observations_text)
            document.add_paragraph("\n")

        if final_shortlist_recommendation and final_shortlist_recommendation.strip() not in ["No final recommendation provided.", ""]:
            document.add_heading("📌 Final Shortlist Recommendation", level=1)
            final_rec_para = document.add_paragraph()
            final_rec_para.add_run(final_shortlist_recommendation).bold = True
            document.add_paragraph("\n")

        doc_io = io.BytesIO()
        document.save(doc_io)
        doc_io.seek(0)
        print("DEBUG (generate_docx_report): DOCX generated successfully.")
        return doc_io
    except Exception as e:
        st.error(f"Error generating DOCX report: {e}")
        print(f"ERROR (generate_docx_report): {e}")
        return None

# --- Supabase Authentication Functions ---
def register_user(email, password, username):
    """Registers a new user in Supabase Auth and stores user profile in the 'users' table."""
    try:
        print(f"DEBUG (register_user): Attempting to sign up user {email}")
        response = supabase.auth.sign_up({"email": email, "password": password})

        if response.user:
            user_id = response.user.id
            # Store user profile in 'users' table using service_role client for direct table access
            # This bypasses RLS for the insert into 'users' table, which is appropriate here
            # as the user is not yet fully authenticated in the app context.
            user_data = {
                'id': user_id,
                'email': email,
                'username': username,
                'created_at': datetime.now().isoformat(), # Use ISO format for Supabase timestamp
                'isadmin': False, # Changed to lowercase 'isadmin'
                'firstloginrequired': True # Changed to lowercase 'firstloginrequired'
            }
            # MODIFIED: Use service_role client for direct table access for user creation
            st.session_state['supabase_service_role_client'].table('users').insert(user_data).execute()
            st.success(f"Account created successfully for {username}! Please check your email to verify and then log in.")
            print(f"DEBUG (register_user): User {username} created and profile saved.")
            time.sleep(2)
            return True
        else:
            st.error(f"Registration failed: {response.session.user.message if response.session and response.session.user else 'Unknown error'}")
            print(f"DEBUG (register_user): Supabase signup failed for {email}. Response: {response}")
            return False
    except Exception as e: # Catching general Exception as APIResponseException import was removed
        error_message = str(e)
        st.error(f"Error creating account: {error_message}")
        print(f"ERROR (register_user): {error_message}")
        time.sleep(2)
        if "User already registered" in error_message or "duplicate key value violates unique constraint" in error_message:
            st.error("This email is already registered.")
        return False

def login_user(email, password, login_as_admin_attempt=False):
    """Logs in a user by verifying their existence in Supabase Auth.
    Also fetches user's admin status from 'users' table and enforces login type.
    Redirects to password update if first login is required."""

    if supabase is None:
        print("ERROR: login_user called but 'supabase' is None. Supabase initialization failed or was not completed.")
        st.error("Application error: Database connection not established. Please refresh or contact support.")
        return False

    try:
        print(f"DEBUG (login_user): Attempting to log in {email}")

        if login_as_admin_attempt:
            # Check against hardcoded admin credentials
            if email == ADMIN_EMAIL and bcrypt.checkpw(password.encode('utf-8'), ADMIN_PASSWORD_HASH.encode('utf-8')):
                st.session_state['logged_in'] = True
                st.session_state['user_email'] = email
                st.session_state['user_name'] = "Admin" # Hardcoded name for the special admin
                st.session_state['user_uid'] = "admin_special_uid" # Placeholder UID for special admin
                st.session_state['is_admin'] = True
                st.success("Admin login successful!")
                print(f"DEBUG (login_user): Special Admin logged in: {email}.")
                time.sleep(1)
                st.session_state['current_page'] = 'Dashboard'
                st.rerun()
                return True
            else:
                st.error("Invalid admin credentials.")
                print(f"DEBUG (login_user): Admin login failed for {email}.")
                return False
        else:
            # Attempt to sign in via Supabase Auth
            response = supabase.auth.sign_in_with_password({"email": email, "password": password})

            if response.user:
                user_id = response.user.id
                # MODIFIED: Fetch user's data from 'users' table using service_role client to bypass RLS for admin check
                user_data_response = st.session_state['supabase_service_role_client'].table('users').select('isadmin', 'firstloginrequired', 'username').eq('id', user_id).single().execute()
                user_data = user_data_response.data if user_data_response.data else {}

                is_user_admin_in_db = user_data.get('isadmin', False) # Changed to lowercase 'isadmin'
                first_login_required = user_data.get('firstloginrequired', True) # Changed to lowercase 'firstloginrequired'
                print(f"DEBUG (login_user): User {email} data: isadmin={is_user_admin_in_db}, firstloginrequired={first_login_required}")

                if login_as_admin_attempt and not is_user_admin_in_db:
                    st.error("This account does not have administrator privileges. Please log in as a regular user.")
                    print(f"DEBUG (login_user): Admin login attempt for non-admin user {email} denied.")
                    return False
                elif not login_as_admin_attempt and is_user_admin_in_db:
                    st.error("This account has administrator privileges. Please log in as an administrator.")
                    print(f"DEBUG (login_user): User login attempt for admin user {email} denied.")
                    return False

                if first_login_required:
                    st.session_state['new_user_email_for_pw_reset'] = email
                    st.session_state['new_user_uid_for_pw_reset'] = user_id
                    st.session_state['current_page'] = 'Update Password'
                    st.success("Please update your password before proceeding.")
                    print(f"DEBUG (login_user): Redirecting {email} to password update page.")
                    time.sleep(1)
                    st.rerun()
                    return True

                st.session_state['logged_in'] = True
                st.session_state['user_email'] = email
                st.session_state['user_name'] = user_data.get('username', email.split('@')[0])
                st.session_state['user_uid'] = user_id
                st.session_state['is_admin'] = is_user_admin_in_db

                st.success(f"Logged in as {st.session_state['user_name']}.")
                if st.session_state['is_admin']:
                    st.info("You are logged in as an administrator.")
                print(f"DEBUG (login_user): Successfully logged in {st.session_state['user_name']} (UID: {st.session_state['user_uid']}, Admin: {st.session_state['is_admin']}).")

                time.sleep(1)
                st.session_state['current_page'] = 'Dashboard'
                st.rerun()
                return True
            else:
                st.error("Invalid email or password.")
                print(f"DEBUG (login_user): Supabase login failed for {email}. Response: {response}")
                return False
    except Exception as e: # Catching general Exception as APIResponseException import was removed
        error_message = str(e)
        st.error(f"An authentication error occurred: {error_message}. Please try again.")
        print(f"ERROR (login_user): Supabase Auth Error during login for {email}: {error_message}")
        time.sleep(2)
        if "Invalid login credentials" in error_message or "Email not confirmed" in error_message:
            st.error("Invalid email or password, or email not confirmed.")
        elif "User not found" in error_message:
            st.error("User not found. Please check your email or sign up.")
        else:
            st.error(f"An unexpected error occurred: {error_message}.")
        return False


def logout_user():
    """Logs out the current user by resetting session state and Supabase session."""
    print("DEBUG (logout_user): Initiating logout.")
    try:
        supabase.auth.sign_out() # Use regular client for logout
        st.session_state['logged_in'] = False
        st.session_state['user_name'] = ''
        st.session_state['user_email'] = ''
        st.session_state['user_uid'] = ''
        st.session_state['is_admin'] = False
        st.session_state['ai_review_result'] = None
        st.session_state['generated_docx_buffer'] = None
        st.session_state['review_triggered'] = False
        st.session_state['current_page'] = 'Login'
        st.session_state['login_mode'] = None
        st.session_state['new_user_email_for_pw_reset'] = ''
        st.session_state['new_user_uid_for_pw_reset'] = ''
        st.success("Logged out successfully!")
        print("DEBUG (logout_user): User logged out. Session state reset. Rerunning.")
        st.rerun()
    except Exception as e:
        st.error(f"Error during logout: {e}")
        print(f"ERROR (logout_user): Error during logout: {e}")

# --- Streamlit Page Functions ---
def dashboard_page():
    """Displays the user dashboard."""
    # Applying color directly with markdown for st.title, as it's not a generic h1 but specific
    st.markdown(f"<h1 style='color: #0D47A1 !important;'>Welcome, {st.session_state['user_name']}!</h1>", unsafe_allow_html=True)
    st.write("This is your dashboard. Use the sidebar to navigate.")
    st.info("To get started, navigate to 'Upload JD & CV' to perform a new AI-powered comparative analysis.")
    if st.session_state['is_admin']: # Only show for admin
        st.write("As an admin, you can also check 'Review Reports' to see all past analyses.")
    print(f"DEBUG (dashboard_page): Displaying dashboard for {st.session_state['user_name']}.") 

def upload_jd_cv_page():
    """Handles JD and CV uploads, triggers AI review, and displays/downloads results."""
    st.markdown("<h1 style='color: #0D47A1 !important;'>⬆️ Upload JD & CV for AI Review</h1>", unsafe_allow_html=True)
    st.write("Upload your Job Description and multiple Candidate CVs to start the comparative analysis.")
    print("DEBUG (upload_jd_cv_page): Displaying upload page.") 

    uploaded_jd = st.file_uploader("Upload Job Description (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], key="jd_uploader")
    uploaded_cvs = st.file_uploader("Upload Candidate's CVs (Multiple - PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True, key="cv_uploader")

    if st.button("Start AI Review", key="start_review_button"):
        print("DEBUG (upload_jd_cv_page): 'Start AI Review' button clicked.") 
        if not uploaded_jd:
            st.warning("Please upload a Job Description.")
            return
        if not uploaded_cvs:
            st.warning("Please upload at least one Candidate CV.")
            return

        jd_text = get_file_content(uploaded_jd, uploaded_jd.name)
        
        all_candidates_data = []
        cv_filenames_list = [] 
        for cv_file in uploaded_cvs:
            cv_text = get_file_content(cv_file, cv_file.name)
            if cv_text:
                all_candidates_data.append({'filename': cv_file.name, 'text': cv_text})
                cv_filenames_list.append(cv_file.name) 
            else:
                st.warning(f"Could not process CV: {cv_file.name}. Skipping it.")
        
        if not jd_text:
            st.error("Failed to extract text from the Job Description.")
            return
        if not all_candidates_data:
            st.error("No valid CVs could be processed for analysis.")
            return

        st.session_state['review_triggered'] = False 
        st.session_state['ai_review_result'] = None
        st.session_state['generated_docx_buffer'] = None

        comparative_results = get_comparative_ai_analysis(jd_text, all_candidates_data)

        if "error" in comparative_results:
            st.error(f"AI analysis failed: {comparative_results['error']}")
        else:
            st.session_state['ai_review_result'] = comparative_results
            st.success("AI review completed successfully!")
            print("DEBUG (upload_jd_cv_page): AI review successful. Preparing DOCX.") 
            
            st.session_state['jd_filename_for_save'] = uploaded_jd.name
            st.session_state['cv_filenames_for_save'] = cv_filenames_list 

            st.session_state['generated_docx_buffer'] = generate_docx_report(
                comparative_results, 
                st.session_state['jd_filename_for_save'], 
                ", ".join(st.session_state['cv_filenames_for_save'])
            )
            
            # --- START OF CHATGPT SUGGESTED CHANGE: SAVE TO CLOUD BEFORE DISPLAYING DOWNLOAD BUTTON ---
            # Generate filename for saving
            timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')
            download_filename = f"{st.session_state['user_name'].replace(' ', '')}_JD-CV_Comparison_Analysis_{timestamp_str}.docx"

            # Adding extra debug prints
            print("DEBUG (upload_jd_cv_page): Calling save_report_on_download now...")
            
            # ✅ CALL save_report_on_download DIRECTLY here
            save_report_on_download(
                download_filename,
                st.session_state['generated_docx_buffer'],
                st.session_state['ai_review_result'],
                st.session_state['jd_filename_for_save'],
                st.session_state['cv_filenames_for_save']
            )
            print("DEBUG (upload_jd_cv_page): save_report_on_download call completed.")
            # --- END OF CHATGPT SUGGESTED CHANGE ---

            st.session_state['review_triggered'] = True 

    if st.session_state['review_triggered'] and st.session_state['ai_review_result']:
        print("DEBUG (upload_jd_cv_page): Displaying AI review results section.") 
        comparative_results = st.session_state['ai_review_result']

        st.subheader("AI Review Results:")

        candidate_evaluations_data = comparative_results.get("candidate_evaluations", [])
        if candidate_evaluations_data:
            st.markdown("### 🧾 Candidate Evaluation Table")
            df_evaluations = pd.DataFrame(candidate_evaluations_data)
            expected_cols_eval = ["Candidate Name", "Match %", "Ranking", "Shortlist Probability", "Key Strengths", "Key Gaps", "Location Suitability", "Comments"]
            for col in expected_cols_eval:
                if col not in df_evaluations.columns:
                    df_evaluations[col] = "N/A"
            df_evaluations = df_evaluations[expected_cols_eval]

            st.dataframe(df_evaluations, use_container_width=True, hide_index=True)
        
        criteria_observations_data = comparative_results.get("criteria_observations", [])
        if criteria_observations_data:
            st.markdown("### ✅ Additional Observations (Criteria Comparison)")
            df_criteria = pd.DataFrame(criteria_observations_data)
            st.dataframe(df_criteria, use_container_width=True, hide_index=True)

        additional_observations_text = comparative_results.get("additional_observations_text", "No general observations provided.")
        if additional_observations_text and additional_observations_text.strip() not in ["No general observations provided.", ""]:
            st.markdown("### General Observations")
            st.write(additional_observations_text)

        final_shortlist_recommendation = comparative_results.get("final_shortlist_recommendation", "No final recommendation provided.")
        if final_shortlist_recommendation and final_shortlist_recommendation.strip() not in ["No final recommendation provided.", ""]:
            st.markdown("### 📌 Final Shortlist Recommendation")
            st.write(final_shortlist_recommendation)

        st.markdown("---") 

        st.subheader("Download & Save Report")
        
        # --- START OF CHATGPT SUGGESTED CHANGE (Modified download button structure) ---
        # No more on_click for save_report_on_download here as it's called earlier
        if st.session_state['generated_docx_buffer']:
            st.download_button(
                label="Download DOCX Report ⬇️", # Label changed for clarity
                data=st.session_state['generated_docx_buffer'],
                file_name=download_filename, # Use the same filename generated for saving
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                key="download_docx_only" # New key for this button
            )
        else:
            st.warning("Run an AI review to generate a report for download and save.")
        # --- END OF CHATGPT SUGGESTED CHANGE ---

# --- Supabase Storage & Database Functions ---

def upload_file_to_supabase(file_bytes, file_name, user_uid):
    """
    Uploads a file to Supabase Storage and returns its public URL.
    Uses service_role client if user_uid is 'admin_special_uid'.
    """
    try:
        bucket_name = "app-files" # Ensure this bucket exists in your Supabase Storage
        file_path_in_storage = f"jd_cv_reports/{user_uid}/{file_name}"

        if user_uid == "admin_special_uid":
            if 'supabase_service_role_client' not in st.session_state:
                st.error("Supabase service role client not initialized.")
                return None
            supabase_target_client = st.session_state['supabase_service_role_client']
            print("DEBUG (upload_file_to_supabase): Using service role client for hardcoded admin upload.")
        else:
            if 'supabase_client' not in st.session_state:
                st.error("Supabase client not initialized.")
                return None
            supabase_target_client = st.session_state['supabase_client']
            print("DEBUG (upload_file_to_supabase): Using regular client for user upload.")

        # Perform the upload.
        # In some versions of supabase-py, this returns an UploadResponse object
        # with 'data' and 'error' as attributes.
        response = supabase_target_client.storage.from_(bucket_name).upload(
            file_path_in_storage,
            file_bytes,
            # It's good practice to specify content-type if known
            # options={"contentType": "application/octet-stream"}
        )

        # Check for error attribute on the response object
        if response.error: # Access 'error' as an attribute
            error_message = response.error.message if hasattr(response.error, 'message') else "Unknown error"
            st.error(f"Supabase Storage upload failed: {error_message}")
            print(f"ERROR (upload_file_to_supabase): Upload failed: {error_message}")
            return None
        else:
            # Upload was successful, now get the public URL
            # The get_public_url method also returns an object, often with a 'data' attribute
            public_url_response = supabase_target_client.storage.from_(bucket_name).get_public_url(file_path_in_storage)

            # Check for error attribute on the public_url_response object
            if public_url_response.error: # Access 'error' as an attribute
                error_message = public_url_response.error.message if hasattr(public_url_response.error, 'message') else "Unknown error"
                st.error(f"Failed to get public URL: {error_message}")
                print(f"ERROR (upload_file_to_supabase): Failed to get public URL: {error_message}")
                return None
            else:
                # The public URL is typically in the 'data' attribute of the public_url_response object
                # and then often as a 'publicUrl' key within that data.
                # Let's be explicit and check if 'data' exists and then if 'publicUrl' is in it.
                if hasattr(public_url_response, 'data') and public_url_response.data:
                    # In some versions, public_url_response.data is directly the URL string.
                    # In others, it's an object like {'publicUrl': '...'}.
                    # We'll try to handle both.
                    if isinstance(public_url_response.data, str):
                        return public_url_response.data
                    elif isinstance(public_url_response.data, dict) and 'publicUrl' in public_url_response.data:
                        return public_url_response.data['publicUrl']
                    else:
                        st.error("Could not extract public URL from Supabase response.")
                        print(f"ERROR (upload_file_to_supabase): Unexpected public URL response data: {public_url_response.data}")
                        return None
                else:
                    st.error("Public URL data not found in Supabase response.")
                    print(f"ERROR (upload_file_to_supabase): Public URL response missing data: {public_url_response}")
                    return None

    except Exception as e:
        st.error(f"An unexpected error occurred during file upload to Supabase Storage: {e}")
        print(f"ERROR (upload_file_to_supabase): Unexpected error: {e}")
        return None

def delete_file_from_supabase_storage(file_path_in_storage, user_uid_for_deletion_check):
    """
    Deletes a file from Supabase Storage.
    Uses service_role client if user_uid_for_deletion_check is 'admin_special_uid'.
    """
    try:
        bucket_name = "app-files" # Ensure this bucket exists in your Supabase Storage

        if user_uid_for_deletion_check == "admin_special_uid":
            if 'supabase_service_role_client' not in st.session_state:
                st.error("Supabase service role client not initialized for deletion.")
                return False
            supabase_target_client = st.session_state['supabase_service_role_client']
            print("DEBUG (delete_file_from_supabase_storage): Using service role client for admin deletion.")
        else:
            if 'supabase_client' not in st.session_state:
                st.error("Supabase client not initialized for deletion.")
                return False
            supabase_target_client = st.session_state['supabase_client']
            print("DEBUG (delete_file_from_supabase_storage): Using regular client for user deletion.")

        # Perform the removal. The remove method expects a list of file paths.
        # It returns an object with 'data' and 'error' as attributes.
        response = supabase_target_client.storage.from_(bucket_name).remove([file_path_in_storage])

        # Check for error attribute on the response object
        if response.error: # Access 'error' as an attribute
            error_message = response.error.message if hasattr(response.error, 'message') else "Unknown error"
            st.error(f"Supabase Storage deletion failed: {error_message}")
            print(f"ERROR (delete_file_from_supabase_storage): Deletion failed: {error_message}")
            return False
        else:
            # If no error, the deletion was successful.
            # The 'data' attribute might contain a list of objects with 'name' and 'id' of deleted files.
            print(f"DEBUG (delete_file_from_supabase_storage): File(s) successfully deleted: {response.data}")
            return True

    except Exception as e:
        st.error(f"An unexpected error occurred during file deletion from Supabase Storage: {e}")
        print(f"ERROR (delete_file_from_supabase_storage): Unexpected error: {e}")
        return False
        
def save_report_on_download(filename, docx_buffer, ai_result, jd_original_name, cv_original_names):
    """Saves the report to Supabase Storage and 'jd_cv_reports' table metadata."""
    st.info("Attempting to save report to cloud... (This message will disappear shortly)")
    print("DEBUG (save_report_on_download): Function started. User UID:", st.session_state.get('user_uid', 'N/A'))

    if supabase is None:
        st.error("Application error: Supabase client not available for saving.")
        print("ERROR (save_report_on_download): Supabase client is None. Cannot save report.")
        return

    storage_file_path = f"jd_cv_reports/{st.session_state['user_uid']}/{filename}"
    download_url = None

    try:
        print(f"DEBUG (save_report_on_download): Attempting to upload file to Storage at: {storage_file_path}")
        docx_buffer.seek(0)
        file_bytes = docx_buffer.getvalue()
        # MODIFIED: Pass user_uid to upload_file_to_supabase to determine client
        download_url = upload_file_to_supabase(file_bytes, filename, st.session_state['user_uid'])

        if download_url:
            st.success(f"File uploaded to Supabase Storage successfully! URL: {download_url}")
            print(f"DEBUG (save_report_on_download): File uploaded to Storage. Public URL: {download_url}")

            report_metadata = {
                "user_email": st.session_state['user_email'],
                "user_name": st.session_state['user_name'],
                "user_uid": st.session_state['user_uid'],
                "jd_filename": jd_original_name,
                "cv_filenames": json.dumps(cv_original_names), # Store list as JSON string
                "review_date": datetime.now().isoformat(), # Use ISO format for Supabase timestamp
                "outputdocfilename": filename, # Changed to lowercase 'outputdocfilename'
                "outputdocurl": download_url, # Changed to lowercase 'outputdocurl'
                "summary": ai_result.get("final_shortlist_recommendation", "No summary provided.")
            }
            print(f"DEBUG (save_report_on_download): Prepared Supabase table metadata: {report_metadata}")

            try:
                print("DEBUG (save_report_on_download): About to attempt saving metadata to Supabase table 'jd_cv_reports'...")
                # ADDED: Use service_role client for hardcoded admin to insert report metadata
                if st.session_state['user_uid'] == "admin_special_uid":
                    print("DEBUG (save_report_on_download): Using service role client for admin metadata insert.")
                    response = st.session_state['supabase_service_role_client'].table('jd_cv_reports').insert(report_metadata).execute()
                else:
                    print("DEBUG (save_report_on_download): Using regular client for user metadata insert.")
                    response = supabase.table('jd_cv_reports').insert(report_metadata).execute()

                if response.data:
                    st.success("Report metadata saved to Supabase successfully!")
                    print("DEBUG (save_report_on_download): Report metadata successfully added to Supabase.")
                else:
                    st.error(f"Supabase metadata save failed: {response.json()}")
                    print(f"ERROR (save_report_on_download): Supabase metadata save failed: {response.json()}")
                    if download_url:
                        # MODIFIED: Pass user_uid to delete_file_from_supabase_storage
                        delete_file_from_supabase_storage(storage_file_path, st.session_state['user_uid'])
                        print("DEBUG: Deleted file from Storage due to metadata save failure.")

            except Exception as generic_e: # Catching general Exception
                st.error(f"An unexpected error occurred during Supabase metadata save: {generic_e}.")
                print(f"ERROR (save_report_on_download): Generic error during Supabase metadata save: {generic_e}")
                if download_url:
                    # MODIFIED: Pass user_uid to delete_file_from_supabase_storage
                    delete_file_from_supabase_storage(storage_file_path, st.session_state['user_uid'])
                    print("DEBUG: Deleted file from Storage due to generic metadata save failure.")
        else:
            st.error("File upload to Supabase Storage failed, so metadata was not saved.")

    except Exception as e: # Catching general Exception
        st.error(f"Error during report upload or initial setup: {e}")
        print(f"ERROR (save_report_on_download): Overall error in function (Storage upload or initial setup): {e}")

    finally:
        pass

def review_reports_page():
    """Displays a table of past reports fetched from Supabase for the current user."""
    st.markdown("<h1 style='color: #0D47A1 !important;'>📚 Review Your Past Reports</h1>", unsafe_allow_html=True)
    st.write("Here you can find a history of your AI-generated comparative analysis reports.")
    print("DEBUG (review_reports_page): Displaying review reports page.")

    if not st.session_state['logged_in'] or not st.session_state['user_uid'] or supabase is None:
        st.info("Please log in to view your past reports.")
        if supabase is None:
            print("ERROR: review_reports_page called but 'supabase' is None.")
        else:
            print("DEBUG (review_reports_page): User not logged in, cannot fetch reports.")
        return

    try:
        print(f"DEBUG (review_reports_page): Fetching reports for UID: {st.session_state['user_uid']}")
        # ADDED: Determine which client to use for fetching reports
        if st.session_state['user_uid'] == "admin_special_uid":
            # Hardcoded admin can view all reports using service_role client
            response = st.session_state['supabase_service_role_client'].table('jd_cv_reports').select('*').execute()
            print("DEBUG (review_reports_page): Admin viewing all reports using service role client.")
        else:
            # Regular user views only their own reports (RLS applies)
            response = supabase.table('jd_cv_reports').select('*').eq('user_uid', st.session_state['user_uid']).execute()
            print("DEBUG (review_reports_page): User viewing own reports using regular client.")

        reviews_data = response.data if response.data else []

        # Sort by review_date in descending order (assuming review_date is ISO format string)
        reviews_data.sort(key=lambda x: x.get('review_date', ''), reverse=True)

        processed_reviews_data = []
        for report in reviews_data:
            cv_filenames = json.loads(report.get('cv_filenames', '[]')) if isinstance(report.get('cv_filenames'), str) else report.get('cv_filenames', [])
            processed_reviews_data.append({
                "Report ID": report.get('id', 'N/A'), # Assuming 'id' is the primary key in Supabase table
                "Report Name": report.get('outputdocfilename', 'N/A'), # Changed to lowercase
                "Job Description": report.get('jd_filename', 'N/A'),
                "Candidates": ", ".join(cv_filenames),
                "Date Generated": datetime.fromisoformat(report['review_date']).strftime('%Y-%m-%d %H:%M:%S') if report.get('review_date') else 'N/A',
                "Summary": report.get('summary', 'No summary provided.'),
                "Download Link": report.get('outputdocurl', '') # Changed to lowercase
            })

        if processed_reviews_data:
            print(f"DEBUG (review_reports_page): Found {len(processed_reviews_data)} reports.")
            df = pd.DataFrame(processed_reviews_data)
            st.dataframe(df,
                         column_config={
                             "Download Link": st.column_config.LinkColumn("Download File", display_text="⬇️ Download", help="Click to download the report file")
                         },
                         hide_index=True,
                         use_container_width=True)
        else:
            st.info("No reports found yet for your account. Start by uploading JD & CVs!")
            print("DEBUG (review_reports_page): No reports found for this user.")
    except Exception as e: # Catching general Exception
        st.error(f"Error fetching your review reports: {e}")
        print(f"ERROR (review_reports_page): Error fetching reports: {e}")


# --- Admin Pages ---
def admin_dashboard_page():
    """Admin dashboard overview."""
    st.markdown("<h1 style='color: #0D47A1 !important;'>⚙️ Admin Dashboard</h1>", unsafe_allow_html=True)
    st.write("Welcome to the Admin Panel. From here you can manage users and all generated reports.")
    st.info("Use the sidebar navigation to access User Management, Report Management, or Invite New Member.")
    print("DEBUG (admin_dashboard_page): Displaying admin dashboard.")

def admin_user_management_page():
    """Admin page to manage users."""
    st.markdown("<h1 style='color: #0D47A1 !important;'>👥 Admin: User Management</h1>", unsafe_allow_html=True)
    st.write("View, manage roles, or delete users.")
    print("DEBUG (admin_user_management_page): Displaying user management page.")

    if st.session_state['supabase_service_role_client'] is None: # MODIFIED: Use service client for admin pages
        print("ERROR: admin_user_management_page called but 'supabase_service_role_client' is None.")
        st.error("Application error: Database connection not established. Please refresh or contact support.")
        return

    users_data = []
    try:
        print(f"DEBUG (admin_user_management_page): Fetching all users from Supabase 'users' table using service role client.")
        # MODIFIED: Use service_role client
        response = st.session_state['supabase_service_role_client'].table('users').select('*').execute()
        users_from_db = response.data if response.data else []

        for user_info in users_from_db:
            users_data.append({
                "UID": user_info.get('id', 'N/A'),
                "Username": user_info.get('username', 'N/A'),
                "Email": user_info.get('email', 'N/A'),
                "Is Admin": user_info.get('isadmin', False) # Changed to lowercase 'isadmin'
            })

        if users_data:
            print(f"DEBUG (admin_user_management_page): Found {len(users_data)} users.")
            df_users = pd.DataFrame(users_data)
            st.dataframe(df_users, use_container_width=True, hide_index=True)

            st.markdown("---")
            st.markdown("<h3 style='color: #0D47A1 !important;'>Manage User Actions</h3>", unsafe_allow_html=True)

            col1, col2 = st.columns(2)

            with col1:
                st.markdown("<h5 style='color: #0D47A1 !important;'>Toggle Admin Status</h5>", unsafe_allow_html=True)
                user_email_toggle = st.text_input("User Email to Toggle Admin", key="toggle_admin_email")
                if st.button("Toggle Admin Status", key="toggle_admin_button"):
                    if user_email_toggle:
                        try:
                            print(f"DEBUG (admin_user_management_page): Toggling admin status for {user_email_toggle}.")
                            # MODIFIED: Fetch user from Supabase Auth to get UID using admin client
                            auth_user_response = st.session_state['supabase_service_role_client'].auth.admin.get_user_by_email(user_email_toggle)
                            user_record = auth_user_response.user

                            if user_record:
                                # MODIFIED: Fetch user data from 'users' table using service client
                                user_data_response = st.session_state['supabase_service_role_client'].table('users').select('isadmin').eq('id', user_record.id).single().execute()
                                user_data = user_data_response.data if user_data_response.data else {}
                                current_admin_status = user_data.get('isadmin', False) # Changed to lowercase 'isadmin'

                                if user_record.id == st.session_state['user_uid'] and current_admin_status:
                                    st.error("You cannot revoke your own administrator privileges.")
                                    print("DEBUG (admin_user_management_page): Self-revocation attempt blocked.")
                                else:
                                    # MODIFIED: Update 'isadmin' in the 'users' table using service client
                                    st.session_state['supabase_service_role_client'].table('users').update({'isadmin': not current_admin_status}).eq('id', user_record.id).execute()
                                    st.success(f"Admin status for {user_email_toggle} toggled to {not current_admin_status}.")
                                    print(f"DEBUG (admin_user_management_page): Admin status for {user_email_toggle} set to {not current_admin_status}.")
                                    time.sleep(1)
                                    st.rerun()
                            else:
                                st.error("User not found in Supabase Auth.")
                                print(f"ERROR (admin_user_management_page): User {user_email_toggle} not found in Supabase Auth.")
                        except Exception as e: # Catching general Exception
                            st.error(f"Error toggling admin status: {e}")
                            print(f"ERROR (admin_user_management_page): Error toggling: {e}")
                    else:
                        st.warning("Please enter a user email to toggle admin status.")

            with col2:
                st.markdown("<h5 style='color: #0D47A1 !important;'>Delete User</h5>", unsafe_allow_html=True)
                user_email_delete = st.text_input("User Email to Delete", key="delete_user_email")
                if st.button("Delete User", key="delete_user_button"):
                    if user_email_delete:
                        if user_email_delete == st.session_state['user_email']:
                            st.error("You cannot delete your own admin account!")
                            print("DEBUG (admin_user_management_page): Self-deletion attempt blocked.")
                        else:
                            try:
                                print(f"DEBUG (admin_user_management_page): Deleting user {user_email_delete}.")
                                # MODIFIED: Get user UID from Supabase Auth using admin client
                                auth_user_response = st.session_state['supabase_service_role_client'].auth.admin.get_user_by_email(user_email_delete)
                                user_record = auth_user_response.user
                                user_uid_to_delete = user_record.id

                                # MODIFIED: Delete associated files from Storage using service client
                                # Need to fetch reports first to get file paths
                                # MODIFIED: Use service_role client and lowercase column name
                                reports_response = st.session_state['supabase_service_role_client'].table('jd_cv_reports').select('outputdocfilename').eq('user_uid', user_uid_to_delete).execute()
                                if reports_response.data:
                                    for report_data in reports_response.data:
                                        storage_file_path = f"jd_cv_reports/{user_uid_to_delete}/{report_data['outputdocfilename']}" # Changed to lowercase
                                        # MODIFIED: Pass "admin_special_uid" to ensure service client is used for deletion
                                        if delete_file_from_supabase_storage(storage_file_path, "admin_special_uid"):
                                            print(f"DEBUG (admin_user_management_page): Deleted Storage file: {storage_file_path}.")
                                        else:
                                            st.warning(f"Could not delete storage file for {user_email_delete}: {report_data['outputdocfilename']}.")

                                # MODIFIED: Delete reports from 'jd_cv_reports' table using service client
                                st.session_state['supabase_service_role_client'].table('jd_cv_reports').delete().eq('user_uid', user_uid_to_delete).execute()
                                print(f"DEBUG (admin_user_management_page): Deleted reports for user {user_email_delete} from 'jd_cv_reports' table.")

                                # MODIFIED: Delete user from 'users' table using service client
                                st.session_state['supabase_service_role_client'].table('users').delete().eq('id', user_uid_to_delete).execute()
                                print(f"DEBUG (admin_user_management_page): Deleted user {user_email_delete} from 'users' table.")

                                # MODIFIED: Delete user from Supabase Auth using admin client
                                st.session_state['supabase_service_role_client'].auth.admin.delete_user(user_uid_to_delete)
                                st.success(f"User {user_email_delete} and all their associated data deleted successfully.")
                                print(f"DEBUG (admin_user_management_page): User {user_email_delete} fully deleted.")
                                time.sleep(1)
                                st.rerun()
                            except Exception as e: # Catching general Exception
                                st.error(f"Error deleting user: {e}")
                                print(f"ERROR (admin_user_management_page): Error deleting user: {e}")
                    else:
                        st.warning("Please enter a user email to delete.")

        else:
            st.info("No users registered yet or error fetching users.")
            print("DEBUG (admin_user_management_page): No users found or fetch error.")

    except Exception as e:
        st.error(f"Error fetching users for admin management: {e}")
        print(f"ERROR (admin_user_management_page): Error fetching users for admin management: {e}")


def admin_report_management_page():
    """Admin page to manage all reports."""
    st.markdown("<h1 style='color: #0D47A1 !important;'>📊 Admin: Report Management</h1>", unsafe_allow_html=True)
    st.write("View and delete all AI-generated comparative analysis reports.")
    print("DEBUG (admin_report_management_page): Displaying report management page.")

    if st.session_state['supabase_service_role_client'] is None: # MODIFIED: Use service client for admin pages
        print("ERROR: admin_report_management_page called but 'supabase_service_role_client' is None.")
        st.error("Application error: Database connection not established. Please refresh or contact support.")
        return

    all_reports_data = []
    try:
        print(f"DEBUG (admin_report_management_page): Fetching all reports from Supabase 'jd_cv_reports' table using service role client.")
        # MODIFIED: Fetching all reports using service_role client
        response = st.session_state['supabase_service_role_client'].table('jd_cv_reports').select('*').execute()
        all_reports_raw = response.data if response.data else []
        all_reports_raw.sort(key=lambda x: x.get('review_date', ''), reverse=True)

        for report_info in all_reports_raw:
            cv_filenames = json.loads(report_info.get('cv_filenames', '[]')) if isinstance(report_info.get('cv_filenames'), str) else report_info.get('cv_filenames', [])
            all_reports_data.append({
                "Report ID": report_info.get('id', 'N/A'),
                "Report Name": report_info.get('outputdocfilename', 'N/A'), # Changed to lowercase
                "Uploaded By": report_info.get('user_name', 'N/A'),
                "Uploader Email": report_info.get('user_email', 'N/A'),
                "JD Filename": report_info.get('jd_filename', 'N/A'),
                "CV Filenames": ", ".join(cv_filenames),
                "Date Generated": datetime.fromisoformat(report_info['review_date']).strftime('%Y-%m-%d %H:%M:%S') if report_info.get('review_date') else 'N/A',
                "Summary": report_info.get('summary', 'No summary provided.'),
                "Download Link": report_info.get('outputdocurl', '') # Changed to lowercase
            })

        if all_reports_data:
            print(f"DEBUG (admin_report_management_page): Found {len(all_reports_data)} reports.")
            df = pd.DataFrame(all_reports_data)
            st.dataframe(df,
                         column_config={
                             "Download Link": st.column_config.LinkColumn("Download File", display_text="⬇️ Download", help="Click to download the report file")
                         },
                         hide_index=True,
                         use_container_width=True)

            st.markdown("---")
            st.markdown("<h3 style='color: #0D47A1 !important;'>Delete Report</h3>", unsafe_allow_html=True)
            report_id_to_delete = st.text_input("Enter Report ID to Delete (from table above)", key="delete_report_id")

            if st.button("Delete Report", key="delete_report_button"):
                if report_id_to_delete:
                    try:
                        print(f"DEBUG (admin_report_management_page): Deleting report {report_id_to_delete}.")
                        # MODIFIED: Fetch report data using service client
                        # MODIFIED: Use service_role client and lowercase column name
                        report_response = st.session_state['supabase_service_role_client'].table('jd_cv_reports').select('outputdocfilename', 'user_uid').eq('id', report_id_to_delete).single().execute()
                        report_data = report_response.data if report_response.data else None

                        if report_data:
                            storage_file_path = f"jd_cv_reports/{report_data['user_uid']}/{report_data['outputdocfilename']}" # Changed to lowercase

                            # MODIFIED: Delete file from storage using service client
                            # MODIFIED: Pass "admin_special_uid" to ensure service client is used for deletion
                            if delete_file_from_supabase_storage(storage_file_path, "admin_special_uid"):
                                st.success(f"File '{report_data['outputdocfilename']}' deleted from Storage.")
                                print(f"DEBUG (admin_report_management_page): Deleted Storage file: {storage_file_path}.")
                            else:
                                st.warning(f"Could not delete storage file for report ID {report_id_to_delete}.")

                            # MODIFIED: Delete report metadata from table using service client
                            response = st.session_state['supabase_service_role_client'].table('jd_cv_reports').delete().eq('id', report_id_to_delete).execute()
                            if response.data:
                                st.success(f"Report '{report_id_to_delete}' deleted from Supabase table.")
                                print(f"DEBUG (admin_report_management_page): Supabase table deletion successful: {report_id_to_delete}.") # Corrected log message
                            else:
                                st.error(f"Failed to delete report from Supabase table: {response.json()}")
                                print(f"ERROR (admin_report_management_page): Supabase table deletion failed: {response.json()}")

                            time.sleep(1)
                            st.rerun()
                        else:
                            st.error("Report with this ID not found.")
                            print(f"ERROR (admin_report_management_page): Report {report_id_to_delete} not found.")
                    except Exception as e: # Catching general Exception
                        st.error(f"Error deleting report: {e}. Ensure Storage path is correct and rules allow deletion.")
                        print(f"ERROR (admin_report_management_page): Error during report deletion for {report_id_to_delete}: {e}")
                else:
                    st.warning("Please enter a Report ID to delete.")

        else:
            st.info("No reports found in the database.")
            print("DEBUG (admin_report_management_page): No reports found in database.")

    except Exception as e: # Catching general Exception
        st.error(f"Error fetching all reports for admin management: {e}")
        print(f"ERROR (admin_report_management_page): Error fetching all reports: {e}")

def admin_invite_member_page():
    """Admin page to invite and create new user accounts."""
    st.markdown("<h1 style='color: #0D47A1 !important;'>➕ Admin: Invite New Member</h1>", unsafe_allow_html=True)
    st.write("Create new user accounts directly and assign their initial role.")
    print("DEBUG (admin_invite_member_page): Displaying invite member page.")

    if st.session_state['supabase_service_role_client'] is None: # MODIFIED: Use service client for admin pages
        print("ERROR: admin_invite_member_page called but 'supabase_service_role_client' is None.")
        st.error("Application error: Database connection not established. Please refresh or contact support.")
        return

    status_message_placeholder = st.empty()

    with st.form("invite_member_form"):
        new_user_email_input = st.text_input("New User Email", help="The email address for the new account.", key="invite_email")
        new_username_input = st.text_input("New User Username", help="A display name for the new user.", key="invite_username")
        new_user_password_input = st.text_input("Temporary Password", type="password", help="A temporary password for the new user. Please communicate this to them securely.", key="invite_password")

        assign_role = st.radio(
            "Assign Role:",
            ("User", "Admin"),
            index=0,
            key="assign_role_radio"
        )

        is_admin_new_user = (assign_role == "Admin")

        confirm_admin_invite = True
        if is_admin_new_user:
            status_message_placeholder.warning("You are about to invite a new Administrator. Administrators have full control over users and reports.")
            confirm_admin_invite = st.checkbox("Yes, I understand and want to create an Administrator account.", key="confirm_admin_invite")

        submit_invite_button = st.form_submit_button("Invite New Member")

        if submit_invite_button:
            print("DEBUG (admin_invite_member_page): 'Invite New Member' button clicked.")
            if is_admin_new_user and not confirm_admin_invite:
                status_message_placeholder.error("Please confirm to create an Administrator account by checking the box.")
                print("DEBUG (admin_invite_member_page): Admin invite: checkbox not confirmed.")
                return

            if not (new_user_email_input and new_username_input and new_user_password_input):
                status_message_placeholder.warning("Please fill in all fields (Email, Username, Temporary Password).")
                print("DEBUG (admin_invite_member_page): Admin invite: missing fields.")
                return

            if not re.match(r"[^@]+@[^@]+\.[^@]+", new_user_email_input):
                status_message_placeholder.warning("Please enter a valid email address.")
                print("DEBUG (admin_invite_member_page): Admin invite: invalid email format.")
                return

            if len(new_user_password_input) < 6:
                status_message_placeholder.warning("Temporary password should be at least 6 characters long (Supabase minimum).")
                print("DEBUG (admin_invite_member_page): Admin invite: weak password.")
                return

            try:
                with st.spinner("Inviting new member..."):
                    print(f"DEBUG (admin_invite_member_page): Attempting to create user {new_user_email_input} with role {assign_role}.")
                    # MODIFIED: Create user in Supabase Auth using service client
                    response = st.session_state['supabase_service_role_client'].auth.admin.create_user(
                        {"email": new_user_email_input, "password": new_user_password_input, "email_confirm": True} # Set email_confirm to True for verification
                    )
                    user_record = response.user

                    if user_record:
                        # MODIFIED: Save user profile in 'users' table using service client
                        user_data = {
                            'id': user_record.id,
                            'email': new_user_email_input,
                            'username': new_username_input,
                            'created_at': datetime.now().isoformat(),
                            'isadmin': is_admin_new_user, # Changed to lowercase
                            'firstloginrequired': True # Changed to lowercase
                        }
                        st.session_state['supabase_service_role_client'].table('users').insert(user_data).execute()

                        status_message_placeholder.success(f"New user '{new_username_input}' ({new_user_email_input}) created successfully with role: {assign_role}! They will need to verify their email.")
                        print(f"DEBUG (admin_invite_member_page): User {new_user_email_input} created in Auth and 'users' table.")

                        st.session_state['invite_email'] = ""
                        st.session_state['invite_username'] = ""
                        st.session_state['invite_password'] = ""
                        st.session_state['assign_role_radio'] = "User"
                        if 'confirm_admin_invite' in st.session_state:
                            st.session_state['confirm_admin_invite'] = False

                        time.sleep(2)
                        st.rerun()
                    else:
                        status_message_placeholder.error(f"Error creating user in Supabase Auth: {response.json()}")
                        print(f"ERROR (admin_invite_member_page): Supabase Auth user creation failed: {response.json()}")

            except Exception as e: # Catching general Exception
                error_message = str(e)
                print(f"ERROR (admin_invite_member_page): Error: {error_message}")
                if "duplicate key value violates unique constraint" in error_message or "email already registered" in error_message:
                    status_message_placeholder.error("This email is already registered. Please use a different email.")
                elif "Password should be at least 6 characters" in error_message:
                    status_message_placeholder.error("The temporary password is too weak. It must be at least 6 characters long.")
                else:
                    status_message_placeholder.error(f"Error inviting new member: {error_message}")

def update_password_page():
    """Page for new users to update their temporary password."""
    st.markdown("<h1 style='color: #0D47A1 !important;'>🔑 Update Your Password</h1>", unsafe_allow_html=True)
    st.write("As a new member, please set your personal password to continue.")
    print("DEBUG (update_password_page): Displaying update password page.")

    if supabase is None:
        print("ERROR: update_password_page called but 'supabase' is None.")
        st.error("Application error: Database connection not established. Please refresh or contact support.")
        return

    if not st.session_state['new_user_uid_for_pw_reset']:
        st.warning("You must be logged in with a temporary account to access this page. Please log in.")
        print("DEBUG (update_password_page): No user UID found for password reset. Redirecting.")
        if st.button("Go to Login"):
            st.session_state['current_page'] = 'Login'
            st.rerun()
        return

    st.info(f"Updating password for: **{st.session_state['new_user_email_for_pw_reset']}**")

    update_status_placeholder = st.empty()

    with st.form("update_password_form"):
        current_temp_password = st.text_input("Current Temporary Password", type="password", help="The password you just used to log in.", key="current_temp_password")
        new_password = st.text_input("New Password", type="password", help="Your new permanent password.", key="new_password_input")
        confirm_new_password = st.text_input("Confirm New Password", type="password", help="Re-enter your new password to confirm.", key="confirm_new_password_input")

        submit_update_button = st.form_submit_button("Update Password")

        if submit_update_button:
            print("DEBUG (update_password_page): 'Update Password' button clicked.")
            if not (current_temp_password and new_password and confirm_new_password):
                update_status_placeholder.warning("Please fill in all password fields.")
                print("DEBUG (update_password_page): Missing password fields.")
                return

            if new_password != confirm_new_password:
                update_status_placeholder.error("New passwords do not match.")
                print("DEBUG (update_password_page): New passwords mismatch.")
                return

            if len(new_password) < 6:
                update_status_placeholder.warning("New password must be at least 6 characters long.")
                print("DEBUG (update_password_page): New password too short.")
                return

            try:
                with st.spinner("Updating password..."):
                    print(f"DEBUG (update_password_page): Attempting to update password for UID: {st.session_state['new_user_uid_for_pw_reset']}")
                    # Re-authenticate with the temporary password to get a valid session for the user
                    auth_response = supabase.auth.sign_in_with_password({
                        "email": st.session_state['new_user_email_for_pw_reset'],
                        "password": current_temp_password
                    })

                    if auth_response.user:
                        # Now that the user is authenticated with the temp password, update their password
                        update_response = supabase.auth.update_user({"password": new_password})
                        if update_response.user:
                            # MODIFIED: Update firstloginrequired status in 'users' table using service client
                            st.session_state['supabase_service_role_client'].table('users').update({'firstloginrequired': False}).eq('id', st.session_state['new_user_uid_for_pw_reset']).execute()

                            update_status_placeholder.success("Password updated successfully! Please log in with your new password.")
                            print("DEBUG (update_password_page): Password updated and firstloginrequired set to False.")
                            time.sleep(2)
                            logout_user() # Log out to force re-login with new password
                        else:
                            update_status_placeholder.error(f"Failed to update password: {update_response.json()}")
                            print(f"ERROR (update_password_page): Supabase update_user failed: {update_response.json()}")
                    else:
                        update_status_placeholder.error("Current temporary password is incorrect.")
                        print(f"ERROR (update_password_page): Invalid temporary password provided for update.")

            except Exception as e: # Catching general Exception
                error_message = str(e)
                print(f"ERROR (update_password_page): Error during password update: {error_message}")
                if "Password should be at least 6 characters" in error_message:
                    update_status_placeholder.error("The new password is too weak. Please choose a stronger one.")
                else:
                    update_status_placeholder.error(f"Error updating password: {error_message}")


# --- Main Streamlit Application Logic ---

def main():
    """Main function to set up Streamlit page and handle navigation/authentication."""
    # Robustly manage current_page state after login/logout
    if st.session_state['logged_in'] and st.session_state['current_page'] in ['Login', 'Signup']:
        st.session_state['current_page'] = 'Dashboard'
    elif not st.session_state['logged_in'] and st.session_state['current_page'] not in ['Login', 'Signup', 'Update Password']:
        st.session_state['current_page'] = 'Login'
        st.session_state['login_mode'] = None

    # Conditional rendering for sidebar (only if logged in)
    if st.session_state['logged_in']:
        with st.sidebar:
            # Sidebar branding
            st.markdown("<h1 style='color: #000000 !important;'>SSO Consultants</h1>", unsafe_allow_html=True) # Forced black
            st.markdown("<h2 style='color: #000000 !important;'>AI Recruitment Dashboard</h2>", unsafe_allow_html=True) # Forced black
            st.markdown("---")

            st.write(f"Welcome, **{st.session_state['user_name']}**!")
            if st.session_state['is_admin']:
                st.markdown("<h3 style='color: #000000 !important;'>Admin Privileges Active</h3>", unsafe_allow_html=True)

            # Navigation for logged-in users (User & Admin)
            user_pages = ['Dashboard', 'Upload JD & CV']
            admin_pages = ['Admin Dashboard', 'Admin: User Management', 'Admin: Report Management', 'Admin: Invite New Member']

            all_pages = user_pages
            if st.session_state['is_admin']:
                all_pages.extend(['Review Reports']) # Add back for admins only
                all_pages.extend(admin_pages)

            try:
                if st.session_state['current_page'] not in all_pages:
                    st.session_state['current_page'] = 'Dashboard'
                default_index = all_pages.index(st.session_state['current_page'])
            except ValueError:
                default_index = 0

            def update_page_selection():
                st.session_state['current_page'] = st.session_state['sidebar_radio_selection']
                print(f"DEBUG (sidebar_radio): Page selected: {st.session_state['current_page']}")

            page_selection = st.radio(
                "Navigation",
                all_pages,
                key="sidebar_radio_selection",
                index=default_index,
                on_change=update_page_selection
            )

            st.markdown("---")
            if st.button("Logout", key="logout_button_sidebar"): # Unique key for sidebar logout
                logout_user()

        # Add a div to the main content area for logged-in users to override centering if needed
        st.markdown('<div class="logged-in-main-content">', unsafe_allow_html=True)

        # --- Render Logged-in Pages ---
        if st.session_state['current_page'] == 'Dashboard':
            dashboard_page()
        elif st.session_state['current_page'] == 'Upload JD & CV':
            upload_jd_cv_page()
        elif st.session_state['is_admin'] and st.session_state['current_page'] == 'Review Reports':
            review_reports_page()
        elif st.session_state['is_admin'] and st.session_state['current_page'] == 'Admin Dashboard':
            admin_dashboard_page()
        elif st.session_state['is_admin'] and st.session_state['current_page'] == 'Admin: User Management':
            admin_user_management_page()
        elif st.session_state['is_admin'] and st.session_state['current_page'] == 'Admin: Report Management':
            admin_report_management_page()
        elif st.session_state['is_admin'] and st.session_state['current_page'] == 'Admin: Invite New Member':
            admin_invite_member_page()
        elif st.session_state['current_page'] == 'Update Password':
             update_password_page()
        else:
            st.error("Access Denied or Page Not Found. Please navigate using the sidebar.")
            print(f"ERROR (main rendering): Invalid page state for logged-in user: {st.session_state['current_page']}")

        st.markdown('</div>', unsafe_allow_html=True) # Close the logged-in-main-content div

    # --- Main Content Area when NOT logged in (Login/Landing Page) ---
    else:
        st.markdown("<h1 class='main-app-title'>SSO Consultants AI Recruitment System</h1>", unsafe_allow_html=True)
        st.markdown("<p class='sub-app-title'>Streamlined Talent Acquisition with AI-Powered Insights</p>", unsafe_allow_html=True)

        # Use columns to center the buttons and potentially the login form
        col_left_spacer_buttons, col_buttons, col_right_spacer_buttons = st.columns([1, 2, 1])
        with col_buttons: # Buttons in the middle column
            admin_col, user_col = st.columns(2)
            with admin_col:
                if st.button("Login as Admin", key="button_login_admin_main_page"):
                    st.session_state['login_mode'] = 'admin'
                    st.session_state['current_page'] = 'Login'
                    print("DEBUG (main): Admin login mode selected from main page.")
                    st.rerun()
            with user_col:
                if st.button("Login as User", key="button_login_user_main_page"):
                    st.session_state['login_mode'] = 'user'
                    st.session_state['current_page'] = 'Login'
                    print("DEBUG (main): User login mode selected from main page.")
                    st.rerun()

        # Display the 'Please select' message
        col_left_spacer_info, col_info_center, col_right_spacer_info = st.columns([1, 2, 1])
        with col_info_center:
            if st.session_state['login_mode'] is None:
                st.markdown("<p class='initial-info-message'>Please select 'Login as Admin' or 'Login as User' to proceed.</p>", unsafe_allow_html=True)

        # Only show login form if a mode has been selected
        if st.session_state['login_mode']:
            col_form_left, col_form_center, col_form_right = st.columns([1, 2, 1])
            with col_form_center: # Form in the middle column
                # The h3 for login form title is explicitly targeted here
                st.markdown(f"<h3 style='text-align: center; color: #000000 !important;'>🔑 Login as {'Administrator' if st.session_state['login_mode'] == 'admin' else 'User'}</h3>", unsafe_allow_html=True)
                with st.form("login_form"):
                    email = st.text_input("Email")
                    password = st.text_input("Password", type="password")
                    submit_button = st.form_submit_button("Login")
                    if submit_button:
                        print(f"DEBUG (main): Login form submitted for {email}.")
                        if email and password:
                            login_user(email, password, login_as_admin_attempt=(st.session_state['login_mode'] == 'admin'))
                        else:
                            st.warning("Please enter both email and password.")

    # --- Custom FOOTER (Always visible at the bottom of the page) ---
    st.markdown(
        """
        <div style="
            position: fixed;
            bottom: 0;
            left: 0;
            width: 100%;
            text-align: center;
            color: #FF8C00 !important; /* Orange text for footer - IMPORTANT for visibility */
            padding: 10px;
            background-color: #FFFFFF; /* Match page background */
            font-size: 0.8em;
            border-top: 1px solid #E0E0E0; /* Subtle border for separation */
            z-index: 999;
        ">
            ©copyright SSO Consultants
        </div>
        """,
        unsafe_allow_html=True
    )

# Entry point for the Streamlit application
if __name__ == "__main__":
    main()
